from __future__ import annotations

import re
from pathlib import Path
from typing import Literal

from django.conf import settings
from django.core.files import File
from django.utils import timezone


def slugify_filename(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r'[^a-z0-9áéíóúãõâêîôûç]+', '-', value)
    value = re.sub(r'-+', '-', value).strip('-')
    return value[:80] or 'material'


def export_content(*, user, title: str, area: str, subject: str, content: str, file_format: Literal['docx', 'pdf', 'txt']) -> Path:
    export_dir = Path(settings.MEDIA_ROOT) / 'exports'
    export_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{slugify_filename(title)}-{timezone.now():%Y%m%d%H%M%S}.{file_format}"
    path = export_dir / filename
    metadata = {
        'user': user.get_full_name() or user.get_username(),
        'title': title,
        'area': area,
        'subject': subject or 'não informada',
        'date': timezone.localtime().strftime('%d/%m/%Y %H:%M'),
    }
    if file_format == 'txt':
        _write_txt(path, metadata, content)
    elif file_format == 'docx':
        _write_docx(path, metadata, content)
    elif file_format == 'pdf':
        _write_pdf(path, metadata, content)
    else:
        raise ValueError('Formato de exportação inválido.')
    return path


def _write_txt(path: Path, metadata: dict, content: str) -> None:
    text = f"""ACADEME.IA

Usuário: {metadata['user']}
Título: {metadata['title']}
Área: {metadata['area']}
Matéria: {metadata['subject']}
Data: {metadata['date']}

{content}

---
Gerado por ACADEME.IA
"""
    path.write_text(text, encoding='utf-8')


def _write_docx(path: Path, metadata: dict, content: str) -> None:
    from docx import Document

    document = Document()
    document.add_heading('ACADEME.IA', level=0)
    document.add_paragraph(f"Usuário: {metadata['user']}")
    document.add_paragraph(f"Título: {metadata['title']}")
    document.add_paragraph(f"Área: {metadata['area']}")
    document.add_paragraph(f"Matéria: {metadata['subject']}")
    document.add_paragraph(f"Data: {metadata['date']}")
    document.add_paragraph('')
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph('')
        elif stripped.startswith('# '):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith('- '):
            document.add_paragraph(stripped[2:], style='List Bullet')
        else:
            document.add_paragraph(stripped)
    document.add_paragraph('Gerado por ACADEME.IA')
    document.save(path)


def _write_pdf(path: Path, metadata: dict, content: str) -> None:
    """Gera um PDF simples sem dependências externas.

    Isso evita problemas de instalação de bibliotecas nativas em versões novas do Python.
    O PDF gerado é textual, suficiente para exportar resumos, transcrições e questões.
    """
    import textwrap

    page_width = 595
    page_height = 842
    left = 54
    top = 790
    line_height = 15
    max_chars = 88

    raw_lines = [
        'ACADEME.IA',
        '',
        f"Usuário: {metadata['user']}",
        f"Título: {metadata['title']}",
        f"Área: {metadata['area']}",
        f"Matéria: {metadata['subject']}",
        f"Data: {metadata['date']}",
        '',
    ]
    raw_lines.extend(content.splitlines())
    raw_lines.extend(['', 'Gerado por ACADEME.IA'])

    wrapped_lines: list[str] = []
    for line in raw_lines:
        clean = line.replace('**', '').replace('#', '').strip()
        if not clean:
            wrapped_lines.append('')
            continue
        wrapped_lines.extend(textwrap.wrap(clean, width=max_chars) or [''])

    pages: list[list[str]] = []
    current: list[str] = []
    lines_per_page = 48
    for line in wrapped_lines:
        if len(current) >= lines_per_page:
            pages.append(current)
            current = []
        current.append(line)
    if current:
        pages.append(current)
    if not pages:
        pages = [['ACADEME.IA']]

    objects: list[bytes] = []

    def pdf_text(value: str) -> bytes:
        value = value.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
        return value.encode('latin-1', errors='replace')

    objects.append(b'<< /Type /Catalog /Pages 2 0 R >>')
    kids = ' '.join(f'{3 + idx * 2} 0 R' for idx in range(len(pages))).encode('ascii')
    objects.append(b'<< /Type /Pages /Kids [' + kids + b'] /Count ' + str(len(pages)).encode('ascii') + b' >>')

    for idx, lines in enumerate(pages):
        page_obj_num = 3 + idx * 2
        content_obj_num = page_obj_num + 1
        page = (
            f'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] '
            f'/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> '
            f'/Contents {content_obj_num} 0 R >>'
        ).encode('ascii')
        objects.append(page)
        stream_parts = [b'BT', b'/F1 11 Tf', f'{left} {top} Td'.encode('ascii')]
        for line_number, line in enumerate(lines):
            if line_number:
                stream_parts.append(f'0 -{line_height} Td'.encode('ascii'))
            stream_parts.append(b'(' + pdf_text(line) + b') Tj')
        stream_parts.append(b'ET')
        stream = b'\n'.join(stream_parts)
        objects.append(b'<< /Length ' + str(len(stream)).encode('ascii') + b' >>\nstream\n' + stream + b'\nendstream')

    output = bytearray(b'%PDF-1.4\n')
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f'{number} 0 obj\n'.encode('ascii'))
        output.extend(obj)
        output.extend(b'\nendobj\n')
    xref_offset = len(output)
    output.extend(f'xref\n0 {len(objects)+1}\n'.encode('ascii'))
    output.extend(b'0000000000 65535 f \n')
    for offset in offsets[1:]:
        output.extend(f'{offset:010d} 00000 n \n'.encode('ascii'))
    output.extend(
        f'trailer\n<< /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n'.encode('ascii')
    )
    path.write_bytes(bytes(output))
